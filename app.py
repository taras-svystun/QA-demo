import streamlit as st
from docx import Document
from io import BytesIO
from langchain.chains import RetrievalQA
from langchain.vectorstores import FAISS
from langchain.document_loaders import TextLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter

from custom_llm import CustomLLM

st.set_page_config(page_title="💬 Chatbot")
st.title("💬 Chatbot for text document QA")

def clear_chat_history():
    st.session_state.messages = [
        {"role": "assistant", "content": "How can I help you?"}
    ]


with st.sidebar:
    st.title('Guidelines 📜')
    st.markdown("""
1. You can ask questions regarding your loaded text files
2. Just add your `.txt` or `.docx` files
                """)
    
    st.title('Notes📌')
    st.markdown("""
1. If you don't have a text file in quick access, the system comes with preloaded test file
2. The user may face 1-2 min delay during the first question due to model loading on the server
                """)

    uploaded_files = st.file_uploader("Add the text files", type=['txt', 'docx'], accept_multiple_files=True)
    st.button(":red[Clear Chat History]", on_click=clear_chat_history)

if bool(uploaded_files):
    content = b''
    for uploaded_file in uploaded_files:
        if uploaded_file.name.endswith("docx"):
            # document = Document()
            i = 0
            for line in uploaded_file:
                # document.add_paragraph(line)
                content += line
                i += 1
                if i >= 5:
                    break
            # content += "".join([paragraph.text for paragraph in document.paragraphs]) + '\n'
            # content += "".join([BytesIO(line) for line in uploaded_file]) + '\n'
            # content = [line for line in uploaded_file]
        else:
            content += "".join([line.decode() for line in uploaded_file]) + '\n'

    st.write(content)
    with open("_sample.txt", "w") as file:
        file.write(content)


loader = TextLoader("_sample.txt") if bool(uploaded_files) else TextLoader("sample.txt")

documents = loader.load()

text_splitter = CharacterTextSplitter(
    separator="\n", chunk_size=1000, chunk_overlap=50
)

text_chunks = text_splitter.split_documents(documents)

texts = text_splitter.split_documents(documents)

embeddings = OpenAIEmbeddings()
vectorstore = FAISS.from_documents(text_chunks, embeddings)

llm = CustomLLM(n=10)

qa = RetrievalQA.from_chain_type(
    llm=llm, chain_type="stuff", retriever=vectorstore.as_retriever()
)


if "messages" not in st.session_state.keys():
    st.session_state.messages = [
        {"role": "assistant", "content": "How can I help you?"}
    ]


for message in st.session_state.messages:
    with st.chat_message(message["role"]):
        st.write(message["content"])


def generate_response(prompt_input):
    output = qa.run(prompt_input)
    return output


if prompt := st.chat_input():
    st.session_state.messages.append({"role": "user", "content": prompt})
    with st.chat_message("user"):
        st.write(prompt)


if st.session_state.messages[-1]["role"] != "assistant":
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            response = generate_response(prompt)
            placeholder = st.empty()
            placeholder.markdown(response)
    message = {"role": "assistant", "content": response}
    st.session_state.messages.append(message)
